import logging
from typing import Optional

logger = logging.getLogger("app.ingestion.parsers.docx")

try:
    import docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.info("python-docx not installed. DOCX parsing will be unavailable.")


class DOCXParser:
    """Extracts text and metadata from DOCX files using python-docx."""

    def parse(self, file_bytes: bytes, filename: str = "") -> Optional["FileParseResult"]:
        from app.ingestion.parsers import FileParseResult

        if not _DOCX_AVAILABLE:
            logger.warning("Cannot parse DOCX: python-docx not installed. Run: pip install python-docx")
            return None

        try:
            import io
            doc = docx.Document(io.BytesIO(file_bytes))
        except Exception as e:
            logger.error(f"Failed to open DOCX '{filename}': {e}")
            return None

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    tables_text.append(" | ".join(cells))

        text_parts = []
        if paragraphs:
            text_parts.append("\n\n".join(paragraphs))
        if tables_text:
            text_parts.append("\n\n=== Tables ===\n" + "\n".join(tables_text))

        text = "\n\n".join(text_parts).strip()

        # Extract metadata from core properties
        core_props = doc.core_properties or None
        title = ""
        author = ""
        if core_props:
            title = (core_props.title or "").strip()
            author = (core_props.author or "").strip()
        if not title:
            title = filename.replace(".docx", "").replace("_", " ").replace("-", " ").title()

        full_metadata = {
            "author": author or "",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
        }
        if core_props:
            full_metadata["created"] = str(core_props.created) if core_props.created else ""

        if not text:
            logger.warning(f"No text extracted from DOCX '{filename}'")
            return None

        logger.info(f"Parsed DOCX '{filename}': {len(paragraphs)} paragraphs, {len(doc.tables)} tables")
        return FileParseResult(
            text=text,
            title=title,
            file_type="docx",
            metadata=full_metadata,
            raw_bytes=file_bytes,
        )


docx_parser = DOCXParser()
